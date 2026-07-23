from fastapi.testclient import TestClient
from io import BytesIO
from docx import Document

from backend.app import app
from backend.services.embeddings import KEYWORD_EMBEDDING_MODE
from conftest import DOCUASK_BACKEND_FAQ_PATH, FAQ_PATH


def test_documents_endpoint_indexes_faq_document():
    client = TestClient(app)
    text = FAQ_PATH.read_text(encoding="utf-8")

    response = client.post(
        "/documents",
        json={
            "text": text,
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "chunk_size": 350,
            "chunk_overlap": 50,
        },
    )

    payload = response.json()

    assert response.status_code == 200
    assert payload["embedding_mode"] == KEYWORD_EMBEDDING_MODE
    assert payload["chunk_count"] == 6
    assert payload["stored_chunk_count"] == 6
    assert payload["collection_name"].startswith("uploaded_document_chunks_keyword_")


def test_documents_upload_endpoint_indexes_markdown_file():
    client = TestClient(app)
    text = FAQ_PATH.read_text(encoding="utf-8")

    response = client.post(
        "/documents/upload",
        data={
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "chunk_size": "350",
            "chunk_overlap": "50",
            },
            files={
                "file": (
                    "rag_faq.md",
                    text.encode("utf-8"),
                    "text/markdown",
                )
        },
    )

    payload = response.json()

    assert response.status_code == 200
    assert payload["embedding_mode"] == KEYWORD_EMBEDDING_MODE
    assert payload["chunk_count"] == 6
    assert payload["stored_chunk_count"] == 6


def test_documents_upload_endpoint_rejects_unsupported_file_type():
    client = TestClient(app)

    response = client.post(
        "/documents/upload",
        data={
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "chunk_size": "350",
            "chunk_overlap": "50",
        },
        files={
            "file": (
                "document.exe",
                b"fake binary",
                "application/octet-stream",
            )
        },
    )

    assert response.status_code == 400
    assert response.json()["detail"]["code"] == "unsupported_file_type"


def test_documents_upload_endpoint_returns_parse_error_for_invalid_pdf():
    client = TestClient(app)

    response = client.post(
        "/documents/upload",
        data={
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "chunk_size": "350",
            "chunk_overlap": "50",
        },
        files={
            "file": (
                "document.pdf",
                b"%PDF-1.4 fake",
                "application/pdf",
            )
        },
    )

    assert response.status_code == 400
    assert response.json()["detail"]["code"] == "document_parse_failed"


def test_documents_upload_endpoint_indexes_docx_file():
    client = TestClient(app)
    buffer = BytesIO()
    document = Document()
    document.add_heading("DocuAsk FAQ", level=1)
    document.add_paragraph("RAG 会先检索资料，再根据资料生成回答。")
    document.save(buffer)

    response = client.post(
        "/documents/upload",
        data={
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "chunk_size": "350",
            "chunk_overlap": "50",
        },
        files={
            "file": (
                "docuask.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    payload = response.json()

    assert response.status_code == 200
    assert payload["chunk_count"] == 1
    assert payload["stored_chunk_count"] == 1


def test_qa_endpoint_retrieves_expected_chunk():
    client = TestClient(app)
    text = FAQ_PATH.read_text(encoding="utf-8")

    document_response = client.post(
        "/documents",
        json={
            "text": text,
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "chunk_size": 350,
            "chunk_overlap": 50,
        },
    )
    collection_name = document_response.json()["collection_name"]

    qa_response = client.post(
        "/qa",
        json={
            "collection_name": collection_name,
            "question": "RAG 的基本流程是什么？",
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "top_k": 3,
        },
    )

    payload = qa_response.json()
    retrieved_chunks = payload["retrieved_chunks"]

    assert qa_response.status_code == 200
    assert payload["retrieval_mode"] == "vector"
    assert 6 in [item["chunk_index"] for item in retrieved_chunks]
    assert len(retrieved_chunks) == 3
    assert "[Chunk 6]" in payload["context"]


def test_qa_endpoint_supports_bm25_mode():
    client = TestClient(app)
    text = FAQ_PATH.read_text(encoding="utf-8")

    document_response = client.post(
        "/documents",
        json={
            "text": text,
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "chunk_size": 350,
            "chunk_overlap": 50,
        },
    )
    collection_name = document_response.json()["collection_name"]

    qa_response = client.post(
        "/qa",
        json={
            "collection_name": collection_name,
            "question": "RAG 的基本流程是什么？",
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "top_k": 3,
            "retrieval_mode": "bm25",
        },
    )

    payload = qa_response.json()
    retrieved_chunks = payload["retrieved_chunks"]

    assert qa_response.status_code == 200
    assert payload["retrieval_mode"] == "bm25"
    assert retrieved_chunks[0]["chunk_index"] == 6
    assert retrieved_chunks[0]["score"] > 0


def test_qa_endpoint_supports_rrf_mode():
    client = TestClient(app)
    text = FAQ_PATH.read_text(encoding="utf-8")

    document_response = client.post(
        "/documents",
        json={
            "text": text,
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "chunk_size": 350,
            "chunk_overlap": 50,
        },
    )
    collection_name = document_response.json()["collection_name"]

    qa_response = client.post(
        "/qa",
        json={
            "collection_name": collection_name,
            "question": "RAG 的基本流程是什么？",
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "top_k": 3,
            "retrieval_mode": "rrf",
        },
    )

    payload = qa_response.json()
    retrieved_chunks = payload["retrieved_chunks"]

    assert qa_response.status_code == 200
    assert payload["retrieval_mode"] == "rrf"
    assert retrieved_chunks[0]["chunk_index"] == 6
    assert retrieved_chunks[0]["rrf_score"] > 0


def test_qa_endpoint_supports_rerank_mode():
    client = TestClient(app)
    text = FAQ_PATH.read_text(encoding="utf-8")

    document_response = client.post(
        "/documents",
        json={
            "text": text,
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "chunk_size": 350,
            "chunk_overlap": 50,
        },
    )
    collection_name = document_response.json()["collection_name"]

    qa_response = client.post(
        "/qa",
        json={
            "collection_name": collection_name,
            "question": "RAG 的基本流程是什么？",
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "top_k": 3,
            "retrieval_mode": "rerank",
        },
    )

    payload = qa_response.json()
    retrieved_chunks = payload["retrieved_chunks"]

    assert qa_response.status_code == 200
    assert payload["retrieval_mode"] == "rerank"
    assert retrieved_chunks[0]["rerank_score"] > 0


def test_qa_endpoint_returns_404_for_missing_collection():
    client = TestClient(app)

    response = client.post(
        "/qa",
        json={
            "collection_name": "missing_collection",
            "question": "RAG 的基本流程是什么？",
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "top_k": 3,
        },
    )

    assert response.status_code == 404


def test_qa_endpoint_returns_400_for_unknown_retrieval_mode():
    client = TestClient(app)
    text = FAQ_PATH.read_text(encoding="utf-8")

    document_response = client.post(
        "/documents",
        json={
            "text": text,
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "chunk_size": 350,
            "chunk_overlap": 50,
        },
    )
    collection_name = document_response.json()["collection_name"]

    response = client.post(
        "/qa",
        json={
            "collection_name": collection_name,
            "question": "RAG 的基本流程是什么？",
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "top_k": 3,
            "retrieval_mode": "unknown",
        },
    )

    assert response.status_code == 400


def test_answer_endpoint_returns_503_when_api_key_is_missing(monkeypatch):
    monkeypatch.delenv("DEEPSEEK_API_KEY", raising=False)

    client = TestClient(app)
    text = FAQ_PATH.read_text(encoding="utf-8")

    document_response = client.post(
        "/documents",
        json={
            "text": text,
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "chunk_size": 350,
            "chunk_overlap": 50,
        },
    )
    collection_name = document_response.json()["collection_name"]

    response = client.post(
        "/answer",
        json={
            "collection_name": collection_name,
            "question": "RAG 的基本流程是什么？",
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "top_k": 3,
            "retrieval_mode": "vector",
        },
    )

    assert response.status_code == 503
    assert response.json()["detail"]["code"] == "missing_api_key"


def test_answer_endpoint_returns_400_for_unknown_retrieval_mode():
    client = TestClient(app)
    text = FAQ_PATH.read_text(encoding="utf-8")

    document_response = client.post(
        "/documents",
        json={
            "text": text,
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "chunk_size": 350,
            "chunk_overlap": 50,
        },
    )
    collection_name = document_response.json()["collection_name"]

    response = client.post(
        "/answer",
        json={
            "collection_name": collection_name,
            "question": "RAG 的基本流程是什么？",
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "top_k": 3,
            "retrieval_mode": "unknown",
        },
    )

    assert response.status_code == 400


def test_evaluation_endpoint_returns_fixed_metrics():
    client = TestClient(app)
    text = FAQ_PATH.read_text(encoding="utf-8")

    response = client.post(
        "/evaluation",
        json={
            "text": text,
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "chunk_size": 350,
            "chunk_overlap": 50,
            "top_k": 3,
        },
    )

    payload = response.json()

    assert response.status_code == 200
    assert payload["retrieval_mode"] == "vector"
    assert payload["chunk_count"] == 6
    assert payload["case_count"] == 15
    assert payload["top_1_hit_rate"] == 0.7333333333333333
    assert payload["top_k_recall"] == 1.0
    assert len(payload["failure_cases"]) >= 1
    rag_flow_row = next(
        row for row in payload["rows"]
        if row["question"] == "RAG 的基本流程是什么？"
    )
    assert rag_flow_row["top_k_hit"] is True


def test_evaluation_endpoint_supports_bm25_mode():
    client = TestClient(app)
    text = FAQ_PATH.read_text(encoding="utf-8")

    response = client.post(
        "/evaluation",
        json={
            "text": text,
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "chunk_size": 350,
            "chunk_overlap": 50,
            "top_k": 3,
            "retrieval_mode": "bm25",
        },
    )

    payload = response.json()

    assert response.status_code == 200
    assert payload["retrieval_mode"] == "bm25"
    assert payload["chunk_count"] == 6
    assert payload["case_count"] == 15
    assert payload["top_k_recall"] >= 0.8


def test_evaluation_endpoint_supports_rrf_mode():
    client = TestClient(app)
    text = FAQ_PATH.read_text(encoding="utf-8")

    response = client.post(
        "/evaluation",
        json={
            "text": text,
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "chunk_size": 350,
            "chunk_overlap": 50,
            "top_k": 3,
            "retrieval_mode": "rrf",
        },
    )

    payload = response.json()

    assert response.status_code == 200
    assert payload["retrieval_mode"] == "rrf"
    assert payload["chunk_count"] == 6
    assert payload["case_count"] == 15
    assert payload["top_k_recall"] >= 0.8


def test_evaluation_endpoint_supports_rerank_mode_and_failure_cases():
    client = TestClient(app)
    text = FAQ_PATH.read_text(encoding="utf-8")

    response = client.post(
        "/evaluation",
        json={
            "text": text,
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "chunk_size": 350,
            "chunk_overlap": 50,
            "top_k": 3,
            "retrieval_mode": "rerank",
        },
    )

    payload = response.json()

    assert response.status_code == 200
    assert payload["retrieval_mode"] == "rerank"
    assert payload["case_count"] == 15
    assert "failure_cases" in payload
    assert all("failure_reason" in row for row in payload["rows"])


def test_evaluation_endpoint_supports_custom_cases():
    client = TestClient(app)
    text = DOCUASK_BACKEND_FAQ_PATH.read_text(encoding="utf-8")

    response = client.post(
        "/evaluation",
        json={
            "text": text,
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "chunk_size": 350,
            "chunk_overlap": 50,
            "top_k": 3,
            "retrieval_mode": "bm25",
            "evaluation_cases": [
                {
                    "question": "FastAPI 后端有哪些接口？",
                    "expected_top_chunk": 2,
                },
                {
                    "question": "documents/upload 文件上传支持什么格式？",
                    "expected_top_chunk": 3,
                },
                {
                    "question": "RRF 混合检索有什么作用？",
                    "expected_top_chunk": 4,
                },
            ],
        },
    )

    payload = response.json()

    assert response.status_code == 200
    assert payload["retrieval_mode"] == "bm25"
    assert payload["chunk_count"] == 4
    assert payload["case_count"] == 3
    assert payload["top_k_recall"] == 1.0


def test_evaluation_endpoint_returns_400_for_unknown_retrieval_mode():
    client = TestClient(app)
    text = FAQ_PATH.read_text(encoding="utf-8")

    response = client.post(
        "/evaluation",
        json={
            "text": text,
            "embedding_mode": KEYWORD_EMBEDDING_MODE,
            "chunk_size": 350,
            "chunk_overlap": 50,
            "top_k": 3,
            "retrieval_mode": "unknown",
        },
    )

    assert response.status_code == 400
